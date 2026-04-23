"""Output resume conversion: Markdown string → .md / .pdf / .docx files."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Protocol

from models.workflow import ResumeTailorResult
from utils.resume_converter import (
    OutputConversionFailedError,
    UnsupportedOutputFormatError,
)


# ---------------------------------------------------------------------------
# Markdown builder
# ---------------------------------------------------------------------------


def build_resume_markdown(result: ResumeTailorResult) -> str:
    """Build a Markdown string from a tailored resume result.

    Args:
        result: ResumeTailorResult with tailored_resume as a JSON-encoded CV dict.

    Returns:
        Formatted Markdown string.
    """
    try:
        cv = json.loads(result.tailored_resume)
    except json.JSONDecodeError as exc:
        raise OutputConversionFailedError(
            f"Invalid JSON in tailored_resume: {exc}"
        ) from exc
    parts: list[str] = [f"# {cv.get('full_name', 'N/A')}\n"]

    if cv.get("contact_info"):
        parts.append(f"{cv['contact_info']}\n")
    parts.append("\n")

    parts.append(f"## Professional Summary\n{cv.get('summary', '')}\n\n")

    parts.append("## Skills\n")
    for skill in cv.get("skills", []):
        parts.append(f"- {skill}\n")

    parts.append("\n## Work Experience\n")
    for exp in cv.get("experience", []):
        parts.append(
            f"### {exp.get('role', '')} at {exp.get('company', '')} "
            f"({exp.get('dates', '')})\n\n"
        )
        for highlight in exp.get("highlights", []):
            parts.append(f"- {highlight}\n")
        parts.append("\n")

    parts.append("## Education\n")
    for edu in cv.get("education", []):
        parts.append(f"- {edu}\n")

    if cv.get("certifications"):
        parts.append("\n## Certifications\n")
        for cert in cv.get("certifications", []):
            parts.append(f"- {cert}\n")

    if cv.get("publications"):
        parts.append("\n## Publications\n")
        for pub in cv.get("publications", []):
            parts.append(f"- {pub}\n")

    if cv.get("projects"):
        parts.append("\n## Projects\n")
        for project in cv.get("projects", []):
            parts.append(f"{project}\n\n")

    return "".join(parts)


# ---------------------------------------------------------------------------
# Protocol
# ---------------------------------------------------------------------------


class OutputConverterProtocol(Protocol):
    """Protocol for output format converters."""

    def convert(self, content: str, output_path: Path) -> None:
        """Write resume content to output_path in the target format."""
        ...


# ---------------------------------------------------------------------------
# Concrete output converters
# ---------------------------------------------------------------------------


class MarkdownOutputConverter:
    """Writes tailored resume as .md."""

    def convert(self, content: str, output_path: Path) -> None:
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(content, encoding="utf-8")
        except OutputConversionFailedError:
            raise
        except Exception as exc:
            raise OutputConversionFailedError(
                f"Failed to write markdown: {exc}"
            ) from exc


class PdfOutputConverter:
    """Writes tailored resume as .pdf — delegates to utils/pdf_converter.py."""

    def convert(self, content: str, output_path: Path) -> None:
        try:
            from utils.pdf_converter import markdown_to_pdf

            output_path.parent.mkdir(parents=True, exist_ok=True)
            markdown_to_pdf(content, str(output_path))
        except OutputConversionFailedError:
            raise
        except Exception as exc:
            raise OutputConversionFailedError(f"Failed to write PDF: {exc}") from exc


class DocxOutputConverter:
    """Writes tailored resume as .docx using python-docx."""

    def convert(self, content: str, output_path: Path) -> None:
        try:
            from docx import Document

            doc = Document()
            for line in content.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:].strip(), style="List Bullet")
                elif line.strip():
                    doc.add_paragraph(line.strip())
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
        except OutputConversionFailedError:
            raise
        except Exception as exc:
            raise OutputConversionFailedError(f"Failed to write docx: {exc}") from exc


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


class OutputConverterRegistry:
    """Maps format strings to output converter implementations."""

    def __init__(self) -> None:
        self._converters: dict[str, OutputConverterProtocol] = {
            "md": MarkdownOutputConverter(),
            "pdf": PdfOutputConverter(),
            "docx": DocxOutputConverter(),
        }

    def get(self, fmt: str) -> OutputConverterProtocol:
        """Return converter for format string.

        Raises:
            UnsupportedOutputFormatError: If fmt is not md, pdf, or docx.
        """
        converter = self._converters.get(fmt.lower())
        if converter is None:
            supported = ", ".join(self._converters)
            raise UnsupportedOutputFormatError(
                f"Unsupported output format '{fmt}'. Supported: {supported}"
            )
        return converter

    def convert_all(
        self, content: str, formats: list[str], output_dir: Path
    ) -> list[Path]:
        """Write content in each requested format.

        Returns:
            List of paths that were written.
        """
        written: list[Path] = []
        for fmt in formats:
            converter = self.get(fmt)
            output_path = output_dir / f"tailored_resume.{fmt}"
            converter.convert(content, output_path)
            written.append(output_path)
        return written
