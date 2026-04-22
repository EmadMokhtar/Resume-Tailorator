import pytest
from pathlib import Path


SAMPLE_MARKDOWN = """\
# Jane Smith

jane@example.com | linkedin.com/in/janesmith

## Professional Summary

Experienced Python engineer.

## Skills

- Python
- Django

## Work Experience

### Senior Engineer at Acme Corp (2020-2024)

- Built microservices

## Education

- BSc CS, State University, 2018
"""


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Jane Smith", 0)
    doc.add_paragraph("jane@example.com | linkedin.com/in/janesmith")
    doc.add_heading("Professional Summary", 1)
    doc.add_paragraph("Experienced Python engineer.")
    doc.add_heading("Skills", 1)
    doc.add_paragraph("Python", style="List Bullet")
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    from markdown_pdf import MarkdownPdf, Section

    pdf = MarkdownPdf()
    pdf.add_section(Section(SAMPLE_MARKDOWN))
    path = tmp_path / "resume.pdf"
    pdf.save(str(path))
    return path


@pytest.fixture
def sample_markdown_path(tmp_path: Path) -> Path:
    path = tmp_path / "resume.md"
    path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")
    return path
